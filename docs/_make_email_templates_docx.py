#!/usr/bin/env python3
"""Render the two CRM email campaign templates
(firebase_cloud/scripts/templates/) into one Word .docx for review.
Re-run after editing the templates."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor

COPPER = RGBColor(0xC6, 0x71, 0x2B)
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TPL = os.path.join(ROOT, "firebase_cloud", "scripts", "templates")

CROWDS = [
    {
        "file": "update_2.2.0.txt",
        "title": "1.  Purchased customers  —  version update",
        "who": "Everyone who bought (status = purchased). Announces what's new in "
               "2.2.0 and how to redownload from their Lemon Squeezy link.",
    },
    {
        "file": "cart_recovery.txt",
        "title": "2.  Cart abandoners  —  recovery",
        "who": "Hit checkout but didn't buy (buyer_lead, not purchased). A nudge to "
               "come back, with a 20% code (REKZ) and a one-line opt-out.",
    },
]


def parse(path):
    lines = open(path, encoding="utf-8").read().split("\n")
    subject, video, i = "", "", 0
    if lines and lines[0].lower().startswith("subject:"):
        subject = lines[0][8:].strip()
        i = 1
    if i < len(lines) and lines[i].lower().startswith("video:"):
        video = lines[i][6:].strip()
        i += 1
    while i < len(lines) and lines[i] == "":
        i += 1
    return subject, video, "\n".join(lines[i:]).strip("\n")


def doc_render(body, video):
    """Show the email the way a recipient sees it: first-name placeholder + the
    video as a labelled link (a clickable thumbnail in the real email)."""
    b = body.replace("{name_part}", " [first name]")
    vid_line = ("▶ Watch the 30-second demo: " + video) if video else ""
    return b.replace("{{video}}", vid_line)


def add_body(doc, body):
    for block in body.split("\n\n"):
        if not block.strip():
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        for i, line in enumerate(block.split("\n")):
            if i:
                p.add_run().add_break()
            run = p.add_run(line)
            run.font.size = Pt(11)
            if line.startswith("▶"):           # the video line, make it pop
                run.bold = True
                run.font.color.rgb = COPPER


doc = Document()
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(doc.sections[0], m, Inches(0.9))

doc.add_heading("Stride — Email Campaigns", level=0)
intro = doc.add_paragraph()
intro.add_run(
    "Two segmented campaigns, both sent via Resend from Joe <home@stridehub.io>. "
    "Each email is personalized with the recipient's first name, and includes the "
    "30-second demo as a clickable video thumbnail (links to "
).font.size = Pt(10)
intro.add_run("youtu.be/CANj6AqznMg").font.size = Pt(10)
intro.add_run(").").font.size = Pt(10)

for crowd in CROWDS:
    subject, video, body = parse(os.path.join(TPL, crowd["file"]))
    body = doc_render(body, video)

    doc.add_heading(crowd["title"], level=1)
    who = doc.add_paragraph()
    r = who.add_run(crowd["who"])
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    subj = doc.add_paragraph()
    subj.add_run("Subject:  ").bold = True
    subj.add_run(subject).font.size = Pt(11)

    add_body(doc, body)

out = os.path.join(ROOT, "docs", "stride-email-campaigns.docx")
doc.save(out)
print("saved", out)
