#!/usr/bin/env python3
"""Build the StrideLink walkthrough VO script as a Word .docx (from the same data
as docs/stridequick-walkthrough-script.md). Re-run after editing the script."""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT

ROWS = [
    (1,  "Intro (title)",            "This is Stride. Shape your whole sound design session from one panel, right inside Ableton.", "7.0s", "0:00 – 0:07"),
    (2,  "Open Canvas",              "Open Canvas brings up the full Stride window where you draw and shape every lane. One press and the engine is ready to go.", "7.0s", "0:07 – 0:14"),
    (3,  "Rescan",                   "Rescan reads your rack and pulls in every parameter you can move. Add a device, hit Rescan, and the new controls land on the lanes.", "7.5s", "0:14 – 0:21.5"),
    (4,  "Lock / Unlock",            "Lock freezes the lanes that are already moving, so your next move only touches the empty ones. Unlock opens them all back up to modulate everything again.", "8.0s", "0:21.5 – 0:29.5"),
    (5,  "Motion row (section intro)", "Every button in the Motion row scans your rack first, then applies its pattern, all in one press. Rescan is optional here, so you never end up with parameters left sitting still.", "7.0s", "0:29.5 – 0:36.5"),
    (6,  "Chaos",                    "Chaos spreads organic, unpredictable motion across every lane at once. Perfect for evolving textures that never sit still.", "7.0s", "0:36.5 – 0:43.5"),
    (7,  "Neuro",                    "Neuro lays down fast, aggressive movement made for bass and modern sound design. One press and the whole rack starts to move.", "7.0s", "0:43.5 – 0:50.5"),
    (8,  "Reflector",                "Reflector mirrors your lanes against each other so they move in symmetry. Pairs rise and fall as reflections for a tight, composed feel.", "7.0s", "0:50.5 – 0:57.5"),
    (9,  "S&H (Sample and Hold)",    "Sample and Hold steps each lane to a new value on the grid. Instant rhythmic, stuttered motion locked to your bars.", "7.0s", "0:57.5 – 1:04.5"),
    (10, "Prism",                    "Prism takes one moving lane and spreads variations of it across all the others. Every parameter follows the same gesture, reshaped.", "7.5s", "1:04.5 – 1:12"),
    (11, "Mutate",                   "Mutate reshapes the curves you already have, chopping them into chunks then shuffling and flipping. It only works on existing motion, so generate or draw a pattern first.", "7.0s", "1:12 – 1:19"),
    (12, "Shuffle",                  "Shuffle trades the motion between lanes so the curves swap places. A fast way to stumble onto a happy accident.", "7.0s", "1:19 – 1:26"),
    (13, "2x / ½x (time stretch)", "2x stretches your motion out, slower and sparser across the bars. Half x compresses it, faster and denser, for double time energy.", "8.0s", "1:26 – 1:34"),
    (14, "Bars (4 / 8 / 16 / 32)",   "Set how many bars your modulation runs for, from 4 up to 32.", "7.5s", "1:34 – 1:41.5"),
    (15, "moving / total (counter)", "This counter shows how many lanes are actually moving out of everything loaded. At a glance you know what still needs a curve.", "7.0s", "1:41.5 – 1:48.5"),
    (16, "Inject to Clip",           "When you're ready to fire your modulation, click Inject to Clip. Every lane writes straight into your Ableton clip, dozens of automation lanes in a single press.", "8.0s", "1:48.5 – 1:56.5"),
    (17, "Outro (CTA)",              "That is Stride. Pick your parameters, shape the motion, inject. stridehub.io", "7.0s", "1:56.5 – 2:03.5"),
]
NOTES = [
    "These VO lines are the captions on screen, so reading them verbatim keeps the audio and captions in sync.",
    "Each beat's length is the dur field for that segment in the reel. Change a dur and the rest re-times automatically.",
    "Brand voice: lead with sound design, producer language, no AI mentions, no em dashes.",
]

doc = Document()
# landscape so the table breathes
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(0.7))

doc.add_heading("StrideLink Walkthrough: VO Script", level=0)
doc.add_paragraph(
    "Voiceover script for the StrideLink walkthrough reel (docs/stridequick-walkthrough.html, "
    "1080x1920, about 2:03 total). Order is top-left to bottom-right across the device. Each line "
    "is the on-screen subtitle, so the VO and captions match. Estimated time is the beat length; "
    "the timecode is where it lands if you read straight through."
)
fmt = doc.add_paragraph()
fmt.add_run("Format: ").bold = True
fmt.add_run("button → description (estimated time)")

table = doc.add_table(rows=1, cols=5)
try:
    table.style = "Light Grid Accent 1"
except KeyError:
    table.style = "Table Grid"
table.autofit = False

HEADERS = ["#", "Button / beat", "Description (VO + subtitle)", "Length", "Timecode"]
for i, h in enumerate(HEADERS):
    c = table.rows[0].cells[i]
    c.text = ""
    run = c.paragraphs[0].add_run(h)
    run.bold = True

for num, button, desc, length, tc in ROWS:
    cells = table.add_row().cells
    cells[0].text = str(num)
    cells[1].text = ""
    cells[1].paragraphs[0].add_run(button).bold = True
    cells[2].text = desc
    cells[3].text = length
    cells[4].text = tc

WIDTHS = [Inches(0.4), Inches(1.7), Inches(5.0), Inches(0.8), Inches(1.3)]
for row in table.rows:
    for i, c in enumerate(row.cells):
        c.width = WIDTHS[i]

tot = doc.add_paragraph()
tot.add_run("Total runtime: 2:03.5").bold = True

doc.add_heading("Notes", level=1)
for n in NOTES:
    doc.add_paragraph(n, style="List Bullet")

out = "docs/stridequick-walkthrough-script.docx"
doc.save(out)
print("saved", out)
